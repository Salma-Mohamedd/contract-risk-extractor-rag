from pathlib import Path
from docx import Document as DocxDocument
from langchain_core.documents import Document


def load_docx(file_path: str) -> list[Document]:
    path = Path(file_path)
    doc = DocxDocument(file_path)

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    text = "\n".join(paragraphs).strip()

    if not text:
        return []

    return [Document(page_content=text, metadata={"source": path.name, "page": 1})]